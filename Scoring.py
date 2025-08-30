import sys
import datetime
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog, filedialog, colorchooser
import json
import os
import math
import statistics

# Gracefully handle the python-docx dependency
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Gracefully handle the matplotlib and numpy dependencies for charting
try:
    import matplotlib.pyplot as plt
    import numpy as np
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False


class ToolTip:
    """Creates a tooltip for a given widget."""
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip_window = None
        self.widget.bind("<Enter>", self.show_tooltip)
        self.widget.bind("<Leave>", self.hide_tooltip)

    def show_tooltip(self, event):
        x, y, _, _ = self.widget.bbox("insert")
        x += self.widget.winfo_rootx() + 25
        y += self.widget.winfo_rooty() + 25

        self.tooltip_window = tk.Toplevel(self.widget)
        self.tooltip_window.wm_overrideredirect(True)
        self.tooltip_window.wm_geometry(f"+{x}+{y}")

        label = tk.Label(self.tooltip_window, text=self.text, justify='left',
                         background="#ffffe0", relief='solid', borderwidth=1,
                         font=("tahoma", "8", "normal"))
        label.pack(ipadx=1)

    def hide_tooltip(self, event):
        if self.tooltip_window:
            self.tooltip_window.destroy()
        self.tooltip_window = None

class ThemeManager:
    """Manages light and dark themes for the entire application."""
    @staticmethod
    def apply_theme(root, theme_name):
        style = ttk.Style(root)
        style.theme_use('clam')

        if theme_name == 'dark':
            # Dark Theme Colors
            bg = '#2e2e2e'
            fg = '#dcdcdc'
            light_bg = '#3c3c3c'
            border_color = '#555555'
            select_bg = '#5a5a5a'

            # Configure base styles
            style.configure('.', background=bg, foreground=fg, bordercolor=border_color, lightcolor=light_bg, darkcolor=bg)
            style.configure('TFrame', background=bg)
            style.configure('TLabel', background=bg, foreground=fg)
            style.configure('TCheckbutton', background=bg, foreground=fg)
            style.configure('TRadiobutton', background=bg, foreground=fg)
            style.configure('TButton', background=light_bg, foreground=fg, bordercolor=border_color)
            style.map('TButton', background=[('active', select_bg)])
            style.configure('TEntry', fieldbackground=light_bg, foreground=fg, insertcolor=fg)
            style.configure('TSpinbox', fieldbackground=light_bg, foreground=fg, insertcolor=fg)
            style.configure('TCombobox', fieldbackground=light_bg, foreground=fg, insertcolor=fg)
            style.configure('TScrollbar', background=bg, troughcolor=light_bg)
            
            # LabelFrame specific style
            style.configure('TLabelFrame', background=bg, bordercolor=border_color)
            style.configure('TLabelFrame.Label', background=bg, foreground=fg)

            # Treeview specific style
            style.configure('Treeview', background=light_bg, fieldbackground=light_bg, foreground=fg)
            style.configure('Treeview.Heading', background=bg, foreground=fg, bordercolor=border_color)
            style.map('Treeview.Heading', background=[('active', select_bg)])
            style.map('Treeview', background=[('selected', '#0078d7')])

            # Apply to root window
            root.config(background=bg)

        else: # Light Theme
            # Revert to default clam styles
            style.configure('.', background='#f0f0f0', foreground='black')
            style.configure('TFrame', background='#f0f0f0')
            style.configure('TLabel', background='#f0f0f0', foreground='black')
            style.configure('TCheckbutton', background='#f0f0f0', foreground='black')
            style.configure('TRadiobutton', background='#f0f0f0', foreground='black')
            style.configure('TButton', background='#f0f0f0', foreground='black')
            style.configure('TEntry', fieldbackground='white', foreground='black', insertcolor='black')
            style.configure('TSpinbox', fieldbackground='white', foreground='black', insertcolor='black')
            style.configure('TCombobox', fieldbackground='white', foreground='black', insertcolor='black')
            style.configure('TScrollbar', background='#f0f0f0', troughcolor='#e1e1e1')
            style.configure('TLabelFrame', background='#f0f0f0')
            style.configure('TLabelFrame.Label', background='#f0f0f0', foreground='black')
            style.configure('Treeview', background='white', fieldbackground='white', foreground='black')
            style.configure('Treeview.Heading', background='#f0f0f0', foreground='black')
            style.map('Treeview', background=[('selected', '#3399ff')])
            root.config(background='#f0f0f0')
        
        # Recursively apply background to all standard tk widgets
        ThemeManager._apply_to_children(root, theme_name)

    @staticmethod
    def _apply_to_children(widget, theme_name):
        bg = '#2e2e2e' if theme_name == 'dark' else '#f0f0f0'
        fg = '#dcdcdc' if theme_name == 'dark' else 'black'
        
        try:
            # Apply to tk widgets that don't use ttk styles
            if 'ttk' not in widget.winfo_class().lower():
                 widget.config(background=bg)
                 if 'label' in widget.winfo_class().lower():
                     widget.config(foreground=fg)
        except tk.TclError:
            pass # Ignore errors for widgets that don't support the config option

        for child in widget.winfo_children():
            ThemeManager._apply_to_children(child, theme_name)

class SettingsWindow(tk.Toplevel):
    """The settings window for configuring scoring rules."""
    def __init__(self, parent, current_settings, app_instance):
        super().__init__(parent)
        self.transient(parent)
        self.title("Settings")
        self.parent = parent
        self.app = app_instance
        self.result = None
        self.grab_set()

        # --- Create local tk variables from current settings ---
        self.settings = current_settings
        self.competition_name_var = tk.StringVar(value=self.settings['competition_name'])
        self.num_categories_var = tk.IntVar(value=self.settings['num_categories'])
        self.enable_curve_var = tk.BooleanVar(value=self.settings['enable_curve'])
        self.enable_weights_var = tk.BooleanVar(value=self.settings['enable_weights'])
        self.score_min_var = tk.StringVar(value=self.settings['score_min'])
        self.score_max_var = tk.StringVar(value=self.settings['score_max'])
        self.dark_mode_var = tk.BooleanVar(value=self.settings.get('dark_mode', False))
        self.chart_type_var = tk.StringVar(value=self.settings.get('chart_type', 'Bar Graph'))

        self.category_names_vars = [tk.StringVar(value=name) for name in self.settings['category_names']]
        self.category_weights_vars = [tk.StringVar(value=weight) for weight in self.settings['category_weights']]
        self.category_calcs_vars = [tk.StringVar(value=calc) for calc in self.settings.get('category_calcs', ['None']*10)]
        self.category_colors_vars = [tk.StringVar(value=color) for color in self.settings.get('category_colors', [])]
        
        # --- UI Creation ---
        self.frame = ttk.Frame(self, padding="15")
        self.frame.pack(expand=True, fill=tk.BOTH)

        # --- General Settings Frame ---
        self.general_frame = ttk.LabelFrame(self.frame, text="General", padding="10")
        self.general_frame.grid(row=0, column=0, sticky="ew", padx=5, pady=5)
        self.general_frame.columnconfigure(1, weight=1)

        ttk.Label(self.general_frame, text="Competition Name:").grid(row=0, column=0, sticky="w")
        ttk.Entry(self.general_frame, textvariable=self.competition_name_var, width=30).grid(row=0, column=1, sticky="ew")

        ttk.Checkbutton(self.general_frame, text="Enable Score Curve", variable=self.enable_curve_var).grid(row=1, column=0, sticky="w", pady=5)
        self.add_tooltip(self.general_frame, row=1, col=2, text="Scales all scores in a category so the top score becomes the max possible score.\nThis proportionally rewards entries that outperform others.")
        
        ttk.Checkbutton(self.general_frame, text="Enable Custom Weights", variable=self.enable_weights_var, command=self.toggle_weight_entries).grid(row=2, column=0, sticky="w")
        self.add_tooltip(self.general_frame, row=2, col=2, text="Allows setting a custom weight for each category's final score contribution.")

        ttk.Label(self.general_frame, text="Score Scale (Min/Max):").grid(row=3, column=0, sticky="w", pady=(10,0))
        scale_frame = ttk.Frame(self.general_frame)
        scale_frame.grid(row=4, column=0, columnspan=2, sticky="w")
        ttk.Entry(scale_frame, textvariable=self.score_min_var, width=5).pack(side=tk.LEFT)
        ttk.Label(scale_frame, text=" to ").pack(side=tk.LEFT)
        ttk.Entry(scale_frame, textvariable=self.score_max_var, width=5).pack(side=tk.LEFT)

        ttk.Label(self.general_frame, text="Report Chart Type:").grid(row=5, column=0, sticky="w", pady=(10,0))
        chart_combo = ttk.Combobox(self.general_frame, textvariable=self.chart_type_var, values=["Bar Graph", "Pie Chart"], state="readonly")
        chart_combo.grid(row=5, column=1, sticky="w", pady=(10,0))
        self.add_tooltip(self.general_frame, row=5, col=2, text="Bar Graph: Compares all submissions.\nPie Chart: Shows a score visual for each entry.")

        # --- Categories Frame ---
        self.categories_outer_frame = ttk.LabelFrame(self.frame, text="Categories", padding="10")
        self.categories_outer_frame.grid(row=1, column=0, sticky="ew", padx=5, pady=5)
        
        ttk.Label(self.categories_outer_frame, text="Number of Categories:").grid(row=0, column=0, sticky="w")
        ttk.Spinbox(self.categories_outer_frame, from_=1, to=10, textvariable=self.num_categories_var, width=5, command=self.rebuild_category_widgets).grid(row=0, column=1, sticky="w")
        
        self.category_frame = ttk.Frame(self.categories_outer_frame, padding=(0, 10, 0, 0))
        self.category_frame.grid(row=1, column=0, columnspan=6, sticky="ew")

        # --- Bottom Frame ---
        bottom_frame = ttk.Frame(self.frame, padding=(0, 10, 0, 0))
        bottom_frame.grid(row=2, column=0, sticky="ew")
        
        ttk.Checkbutton(bottom_frame, text="Dark Mode", variable=self.dark_mode_var, command=self.toggle_dark_mode).pack(side=tk.LEFT, padx=5)
        
        credit_frame = ttk.Frame(bottom_frame)
        credit_frame.pack(side=tk.LEFT, expand=True, padx=10)
        self.credit_label = ttk.Label(credit_frame, text="Developed by Noah Hull", font=("Helvetica", 7))
        self.credit_label.pack(side=tk.TOP, anchor='w')
        self.email_label = ttk.Label(credit_frame, text="noahhull1999@gmail.com", font=("Helvetica", 7))
        self.email_label.pack(side=tk.TOP, anchor='w')

        ttk.Button(bottom_frame, text="Save", command=self.on_save).pack(side=tk.RIGHT, padx=5)
        ttk.Button(bottom_frame, text="Cancel", command=self.destroy).pack(side=tk.RIGHT)

        self.rebuild_category_widgets()
        self.app.update_theme() # Apply initial theme

    def add_tooltip(self, parent, row, col, text):
        info_label = ttk.Label(parent, text=" (?)", cursor="question_arrow")
        info_label.grid(row=row, column=col, sticky="w", padx=2)
        ToolTip(info_label, text)

    def _choose_color(self, index, swatch_label):
        """Opens a color chooser and updates the variable and swatch."""
        color_code = colorchooser.askcolor(title="Choose category color")
        if color_code and color_code[1]:
            self.category_colors_vars[index].set(color_code[1])
            swatch_label.config(background=color_code[1])

    def rebuild_category_widgets(self):
        for widget in self.category_frame.winfo_children():
            widget.destroy()

        num = self.num_categories_var.get()
        
        # Ensure settings lists are long enough
        default_colors = self.app.default_settings['category_colors']
        while len(self.category_names_vars) < num: self.category_names_vars.append(tk.StringVar(value=f"Category {len(self.category_names_vars)+1}"))
        while len(self.category_weights_vars) < num: self.category_weights_vars.append(tk.StringVar(value=str(round(100/num))))
        while len(self.category_calcs_vars) < num: self.category_calcs_vars.append(tk.StringVar(value='None'))
        while len(self.category_colors_vars) < num: self.category_colors_vars.append(tk.StringVar(value=default_colors[len(self.category_colors_vars)]))

        # --- Create Headers ---
        ttk.Label(self.category_frame, text="Category Name", font="-weight bold").grid(row=0, column=1, padx=5, pady=5)
        ttk.Label(self.category_frame, text="Weight (%)", font="-weight bold").grid(row=0, column=2, padx=5, pady=5)
        calc_header = ttk.Label(self.category_frame, text="Pre-Calculation", font="-weight bold")
        calc_header.grid(row=0, column=3, padx=5, pady=5)
        ttk.Label(self.category_frame, text="Chart Color", font="-weight bold").grid(row=0, column=4, columnspan=2, padx=5, pady=5)
        
        tooltip_text = """Apply a function to scores before other calculations.

--- Basic Adjustments ---
• Square Root (√x): Reduces the impact of very high scores.
• Log10 (log₁₀x): Aggressively reduces impact of high scores.
• Square (x²): Magnifies the difference between high scores.
• Invert (Max - x): For "lower is better" scores (e.g., time).
• Binary (Pass/Fail): Any score > 0 becomes the max score.

--- Relational Adjustments (compares all submissions) ---
• Z-Score: Normalizes score based on category average & std dev.
  Excellent for rewarding standout performance.
• Rank Order: Score is based on rank (1st, 2nd, 3rd) in category.
• Diff from Average: Score becomes the raw score minus the average.
• Pct of Top Score: Score is its percentage of the top raw score."""
        self.add_tooltip(self.category_frame, row=0, col=6, text=tooltip_text)

        self.weight_entries = []
        calc_options = ["None", "Square Root", "Log10", "Square", "Invert (Max - x)", "Binary (Pass/Fail)",
                        "Z-Score", "Rank Order", "Diff from Average", "Pct of Top Score"]
        for i in range(num):
            ttk.Label(self.category_frame, text=f"Category {i+1}:").grid(row=i+1, column=0, sticky="w")
            ttk.Entry(self.category_frame, textvariable=self.category_names_vars[i]).grid(row=i+1, column=1, padx=5, pady=2, sticky="ew")
            
            weight_entry = ttk.Entry(self.category_frame, textvariable=self.category_weights_vars[i], width=10)
            weight_entry.grid(row=i+1, column=2, padx=5, pady=2)
            self.weight_entries.append(weight_entry)
            
            calc_combo = ttk.Combobox(self.category_frame, textvariable=self.category_calcs_vars[i], values=calc_options, state="readonly", width=18)
            calc_combo.grid(row=i+1, column=3, padx=5, pady=2)

            # Color Swatch & Button
            color_swatch = tk.Label(self.category_frame, text="    ", background=self.category_colors_vars[i].get(), relief="sunken")
            color_swatch.grid(row=i+1, column=4, padx=(5,2), pady=2, sticky="ew")
            
            color_button = ttk.Button(self.category_frame, text="Choose", command=lambda idx=i, swatch=color_swatch: self._choose_color(idx, swatch))
            color_button.grid(row=i+1, column=5, padx=(0,5), pady=2)

        self.toggle_weight_entries()
        self.app.update_theme()

    def toggle_weight_entries(self):
        state = tk.NORMAL if self.enable_weights_var.get() else tk.DISABLED
        for entry in self.weight_entries:
            entry.config(state=state)

    def toggle_dark_mode(self):
        self.app.settings['dark_mode'] = self.dark_mode_var.get()
        self.app.update_theme()
        ThemeManager.apply_theme(self, 'dark' if self.app.settings['dark_mode'] else 'light')


    def on_save(self):
        num_cat = self.num_categories_var.get()
        if self.enable_weights_var.get():
            try:
                weights = [int(self.category_weights_vars[i].get() or 0) for i in range(num_cat)]
                if sum(weights) != 100:
                    messagebox.showerror("Validation Error", f"Weights must add up to 100. Current sum is {sum(weights)}.", parent=self)
                    return
            except ValueError:
                messagebox.showerror("Validation Error", "All weight fields must be valid integers.", parent=self)
                return
        try:
            int(self.score_min_var.get())
            int(self.score_max_var.get())
        except ValueError:
            messagebox.showerror("Validation Error", "Score scale values must be integers.", parent=self)
            return

        self.result = {
            'competition_name': self.competition_name_var.get().strip() or "Untitled Competition",
            'num_categories': num_cat,
            'enable_curve': self.enable_curve_var.get(),
            'enable_weights': self.enable_weights_var.get(),
            'score_min': self.score_min_var.get(),
            'score_max': self.score_max_var.get(),
            'dark_mode': self.dark_mode_var.get(),
            'chart_type': self.chart_type_var.get(),
            'category_names': [var.get() or f"Category {i+1}" for i, var in enumerate(self.category_names_vars[:num_cat])],
            'category_weights': [var.get() or str(round(100/num_cat)) for i, var in enumerate(self.category_weights_vars[:num_cat])],
            'category_calcs': [var.get() for var in self.category_calcs_vars[:num_cat]],
            'category_colors': [var.get() for var in self.category_colors_vars[:num_cat]]
        }
        self.destroy()

class ScoreCalculatorApp:
    CONFIG_FILE = "config.json"

    def __init__(self, root):
        self.root = root
        self.root.geometry("1100x600")
        self.active_windows = [self.root]

        self.default_settings = {
            'competition_name': "New Competition", 'num_categories': 3,
            'category_names': ["Combat", "Design", "Creativity"], 
            'category_weights': ["60", "20", "20"],
            'category_calcs': ["None", "None", "None"],
            'enable_curve': True, 'enable_weights': True,
            'score_min': "1", 'score_max': "100", 'dark_mode': False,
            'chart_type': 'Bar Graph',
            'category_colors': ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', 
                                '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf']
        }
        self.settings = self.default_settings.copy()
        self.load_config()

        self.submissions = []
        self.selected_item_id = None
        
        self.create_widgets()
        self.apply_settings_to_ui()
        self.update_theme()
        
        self.root.protocol("WM_DELETE_WINDOW", self.save_config_on_exit)

    def load_config(self):
        """Loads settings from the config file on startup."""
        try:
            with open(self.CONFIG_FILE, 'r') as f:
                config_data = json.load(f)
                # Ensure all default keys are present after loading
                for key, value in self.default_settings.items():
                    config_data.setdefault(key, value)
                self.settings.update(config_data)
        except (FileNotFoundError, json.JSONDecodeError):
            # If file doesn't exist or is corrupt, just use defaults
            pass

    def save_config_on_exit(self):
        """Saves current settings to the config file and closes the app."""
        try:
            with open(self.CONFIG_FILE, 'w') as f:
                json.dump(self.settings, f, indent=4)
        except Exception as e:
            # Log error or inform user, but don't prevent closing
            print(f"Error saving config: {e}")
        self.root.destroy()

    def update_theme(self):
        theme = 'dark' if self.settings.get('dark_mode', False) else 'light'
        for window in self.active_windows:
            if window.winfo_exists():
                ThemeManager.apply_theme(window, theme)
            else:
                self.active_windows.remove(window)

    def create_widgets(self):
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New Competition", command=self.new_competition)
        file_menu.add_command(label="Load Competition...", command=self.load_competition)
        file_menu.add_command(label="Save Competition As...", command=self.save_competition)
        file_menu.add_separator()
        file_menu.add_command(label="Settings...", command=self.open_settings_window)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.save_config_on_exit)

        self.main_frame = ttk.Frame(self.root, padding="10")
        self.main_frame.pack(fill=tk.BOTH, expand=True)
        
        self.main_frame.grid_rowconfigure(0, weight=1)
        self.main_frame.grid_columnconfigure(0, weight=1)
        self.main_frame.grid_columnconfigure(1, weight=0)

        left_frame = ttk.Frame(self.main_frame)
        left_frame.grid(row=0, column=0, sticky='nsew', padx=(0, 10))

        self.right_frame = ttk.LabelFrame(self.main_frame, text="Submission Details", padding="10")
        self.right_frame.grid(row=0, column=1, sticky='ns')

        self.tree = ttk.Treeview(left_frame, show="headings", selectmode="browse")
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar = ttk.Scrollbar(left_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.bind("<<TreeviewSelect>>", self.on_submission_select)
        
        bottom_buttons_frame = ttk.Frame(self.main_frame)
        bottom_buttons_frame.grid(row=1, column=0, columnspan=2, sticky='ew', pady=(10,0))
        bottom_buttons_frame.grid_columnconfigure(0, weight=1)
        bottom_buttons_frame.grid_columnconfigure(1, weight=1)

        self.report_button = ttk.Button(bottom_buttons_frame, text="📊 Generate Report", command=self.generate_report)
        self.report_button.grid(row=0, column=0, sticky='ew', padx=(0,5))
        self.settings_button = ttk.Button(bottom_buttons_frame, text="⚙️ Settings", command=self.open_settings_window)
        self.settings_button.grid(row=0, column=1, sticky='ew', padx=(5,0))
    
    def apply_settings_to_ui(self):
        self.root.title(f"{self.settings['competition_name']} - Score Calculator")
        
        self.tree.delete(*self.tree.get_children())
        cols = ("name",) + tuple(f"cat_{i}" for i in range(self.settings['num_categories']))
        self.tree["columns"] = cols
        self.tree.heading("name", text="Submission Name")
        self.tree.column("name", width=180, minwidth=150)
        for i, name in enumerate(self.settings['category_names']):
            self.tree.heading(f"cat_{i}", text=name)
            self.tree.column(f"cat_{i}", width=60, anchor=tk.CENTER)

        for widget in self.right_frame.winfo_children():
            widget.destroy()
        
        self.right_frame.grid_rowconfigure(99, weight=1)
        self.right_frame.grid_columnconfigure(0, weight=1)

        ttk.Label(self.right_frame, text="Submission Name:").grid(row=0, column=0, columnspan=2, sticky="w", pady=(0,2))
        self.name_var = tk.StringVar()
        self.name_entry = ttk.Entry(self.right_frame, textvariable=self.name_var, width=35)
        self.name_entry.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0,10))
        
        self.detail_score_vars = []
        for i, name in enumerate(self.settings['category_names']):
            ttk.Label(self.right_frame, text=f"{name} Score:").grid(row=i*2+2, column=0, columnspan=2, sticky="w", pady=(0,2))
            var = tk.StringVar()
            ttk.Entry(self.right_frame, textvariable=var).grid(row=i*2+3, column=0, columnspan=2, sticky="ew", pady=(0,10))
            self.detail_score_vars.append(var)

        notes_y_pos = self.settings['num_categories'] * 2 + 4
        ttk.Label(self.right_frame, text="Notes:").grid(row=notes_y_pos, column=0, columnspan=2, sticky="w", pady=(5,2))
        notes_frame = ttk.Frame(self.right_frame)
        notes_frame.grid(row=notes_y_pos + 1, column=0, columnspan=2, sticky="nsew", pady=(0,10))
        notes_frame.grid_rowconfigure(0, weight=1)
        notes_frame.grid_columnconfigure(0, weight=1)
        self.notes_text = tk.Text(notes_frame, height=5, wrap=tk.WORD, relief=tk.FLAT)
        self.notes_text.grid(row=0, column=0, sticky="nsew")
        notes_scroll = ttk.Scrollbar(notes_frame, orient=tk.VERTICAL, command=self.notes_text.yview)
        notes_scroll.grid(row=0, column=1, sticky="ns")
        self.notes_text['yscrollcommand'] = notes_scroll.set
        
        button_y_pos = notes_y_pos + 2
        ttk.Button(self.right_frame, text="💾 Save", command=self.save_submission).grid(row=button_y_pos, column=0, columnspan=2, sticky="ew", pady=5)
        ttk.Button(self.right_frame, text="➕ Add New", command=self.clear_form).grid(row=button_y_pos+1, column=0, sticky="ew", padx=(0,2))
        ttk.Button(self.right_frame, text="🗑️ Delete Selected", command=self.delete_submission).grid(row=button_y_pos+1, column=1, sticky="ew", padx=(2,0))
        self.update_theme()
    
    def open_settings_window(self):
        settings_win = SettingsWindow(self.root, self.settings, self)
        self.active_windows.append(settings_win)
        self.root.wait_window(settings_win)
        
        if settings_win.result:
            if self.settings['num_categories'] != settings_win.result['num_categories'] and self.submissions:
                if messagebox.askyesno("Confirm Change", "Changing the number of categories requires clearing all current submissions. Proceed?", parent=self.root):
                    self.submissions = []
                else:
                    return
            self.settings.update(settings_win.result)
            self.apply_settings_to_ui()
            self.refresh_submission_list()
            self.update_theme()

    def refresh_submission_list(self):
        self.tree.delete(*self.tree.get_children())
        for i, sub in enumerate(self.submissions):
            values = (sub['name'],) + tuple(sub['scores'])
            self.tree.insert("", tk.END, iid=i, values=values)

    def on_submission_select(self, event):
        selected_items = self.tree.selection()
        if not selected_items: return
        self.selected_item_id = int(selected_items[0])
        submission = self.submissions[self.selected_item_id]
        self.name_var.set(submission['name'])
        for i, score in enumerate(submission['scores']):
            self.detail_score_vars[i].set(score)
        self.notes_text.delete("1.0", tk.END)
        self.notes_text.insert("1.0", submission.get('notes', ''))

    def clear_form(self):
        self.selected_item_id = None
        self.name_var.set("")
        for var in self.detail_score_vars: var.set("")
        self.notes_text.delete("1.0", tk.END)
        self.tree.selection_remove(self.tree.selection())
        self.name_entry.focus()

    def save_submission(self):
        name = self.name_var.get().strip()
        if not name:
            messagebox.showerror("Validation Error", "Submission Name cannot be empty.", parent=self.root)
            return
        
        scores = [var.get().strip() for var in self.detail_score_vars]
        notes = self.notes_text.get("1.0", tk.END).strip()
        submission_data = {'name': name, 'scores': scores, 'notes': notes}

        if self.selected_item_id is not None:
            self.submissions[self.selected_item_id] = submission_data
        else:
            self.submissions.append(submission_data)

        self.refresh_submission_list()
        self.clear_form()

    def delete_submission(self):
        if self.selected_item_id is None:
            messagebox.showwarning("Selection Error", "Please select a submission to delete.", parent=self.root)
            return
        
        name = self.submissions[self.selected_item_id]['name']
        if messagebox.askyesno("Confirm Delete", f"Are you sure you want to delete '{name}'?", parent=self.root):
            del self.submissions[self.selected_item_id]
            self.refresh_submission_list()
            self.clear_form()
    
    def new_competition(self):
        if messagebox.askyesno("Confirm", "Create a new competition? Any unsaved changes will be lost.", parent=self.root):
            self.submissions = []
            self.settings = self.default_settings.copy()
            self.apply_settings_to_ui()
            self.refresh_submission_list()
            self.update_theme()

    def save_competition(self):
        save_dir = "competitions"
        os.makedirs(save_dir, exist_ok=True)
        
        filepath = filedialog.asksaveasfilename(
            initialdir=save_dir,
            title="Save Competition",
            defaultextension=".json",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
        )
        if not filepath: return
        
        data_to_save = {
            'settings': self.settings,
            'submissions': self.submissions
        }
        try:
            with open(filepath, 'w') as f:
                json.dump(data_to_save, f, indent=4)
            messagebox.showinfo("Success", f"Competition saved successfully to:\n{os.path.abspath(filepath)}", parent=self.root)
        except Exception as e:
            messagebox.showerror("Save Error", f"An error occurred while saving:\n{e}", parent=self.root)

    def load_competition(self):
        if messagebox.askyesno("Confirm", "Load a competition? Any unsaved changes will be lost.", parent=self.root):
            filepath = filedialog.askopenfilename(
                initialdir="competitions",
                title="Load Competition",
                filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
            )
            if not filepath: return

            try:
                with open(filepath, 'r') as f:
                    data = json.load(f)
                temp_settings = self.default_settings.copy()
                temp_settings.update(data['settings'])
                self.settings = temp_settings
                self.submissions = data['submissions']
                self.apply_settings_to_ui()
                self.refresh_submission_list()
                self.update_theme()
            except Exception as e:
                messagebox.showerror("Load Error", f"An error occurred while loading:\n{e}", parent=self.root)
    
    def _apply_pre_calc(self, score, calc_type, stats, sub_name):
        if calc_type == "None": return score
        if score <= 0 and calc_type in ["Square Root", "Log10"]: return 0
        
        max_val = int(self.settings.get('score_max', 100))

        if calc_type == "Square Root": return math.sqrt(score)
        if calc_type == "Log10": return math.log10(score) * (max_val / 2)
        if calc_type == "Square": return score ** 2
        if calc_type == "Invert (Max - x)": return max(0, max_val - score)
        if calc_type == "Binary (Pass/Fail)": return max_val if score > 0 else 0
        
        if calc_type == "Z-Score":
            return ((score - stats['avg']) / stats['std']) * 10 + 50 if stats['std'] > 0 else 50
        if calc_type == "Rank Order":
            return stats['rank_map'].get(sub_name, 0)
        if calc_type == "Diff from Average":
            return score - stats['avg']
        if calc_type == "Pct of Top Score":
            return (score / stats['max']) * max_val if stats['max'] > 0 else 0
            
        return score

    def generate_report(self):
        if not self.submissions:
            messagebox.showinfo("No Data", "There are no submissions to report.", parent=self.root)
            return

        num_cat = self.settings['num_categories']
        calcs = self.settings.get('category_calcs', ['None'] * num_cat)
        
        raw_scores_by_cat = [[] for _ in range(num_cat)]
        validated_subs = []
        for sub in self.submissions:
            try:
                raw_scores = [float(s or 0) for s in sub['scores']]
                validated_subs.append({'name': sub['name'], 'raw_scores': raw_scores})
                for i in range(num_cat):
                    raw_scores_by_cat[i].append(raw_scores[i])
            except (ValueError, IndexError):
                messagebox.showerror("Score Error", f"Submission '{sub['name']}' has an invalid score. Please correct it.", parent=self.root)
                return

        category_stats = []
        for i in range(num_cat):
            scores = raw_scores_by_cat[i]
            stats = {}
            if len(scores) > 1:
                stats['avg'] = statistics.mean(scores)
                stats['std'] = statistics.stdev(scores)
            else:
                stats['avg'] = scores[0] if scores else 0
                stats['std'] = 0
            stats['max'] = max(scores) if scores else 0
            stats['min'] = min(scores) if scores else 0
            
            sorted_subs = sorted(validated_subs, key=lambda s: s['raw_scores'][i], reverse=True)
            max_rank_score = len(sorted_subs)
            stats['rank_map'] = {sub['name']: max_rank_score - rank for rank, sub in enumerate(sorted_subs)}
            
            category_stats.append(stats)
            
        processed_subs = []
        for sub in validated_subs:
            pre_calc_scores = [self._apply_pre_calc(s, calcs[i], category_stats[i], sub['name']) for i, s in enumerate(sub['raw_scores'])]
            processed_subs.append({'name': sub['name'], 'raw_scores': sub['raw_scores'], 'pre_calc_scores': pre_calc_scores})
        
        max_pre_calc_scores = [max((sub['pre_calc_scores'][i] for sub in processed_subs), default=1) for i in range(num_cat)]
        min_pre_calc_scores = [min((sub['pre_calc_scores'][i] for sub in processed_subs), default=0) for i in range(num_cat)]

        results = []
        for sub in processed_subs:
            base_scores = sub['pre_calc_scores']
            final_category_scores = base_scores
            if self.settings['enable_curve']:
                final_category_scores = []
                try:
                    score_max_value = float(self.settings.get('score_max', 100.0))
                except (ValueError, TypeError):
                    score_max_value = 100.0
                
                for i in range(num_cat):
                    score = base_scores[i]
                    max_s = max_pre_calc_scores[i]
                    min_s = min_pre_calc_scores[i]

                    if max_s > 0:
                        scaling_factor = score_max_value / max_s
                        curved_score = score * scaling_factor
                    else:
                        if max_s == min_s:
                            curved_score = score_max_value
                        else:
                            curved_score = score_max_value * (score - min_s) / (max_s - min_s)
                    
                    final_category_scores.append(curved_score)

            final_score = sum(final_category_scores) / num_cat if num_cat > 0 else 0
            if self.settings['enable_weights']:
                weights = [int(w) for w in self.settings['category_weights']]
                final_score = sum(s * (w / 100) for s, w in zip(final_category_scores, weights))

            results.append({'name': sub['name'], 'raw_scores': sub['raw_scores'], 'final_category_scores': final_category_scores, 'final_score': final_score})

        sorted_results = sorted(results, key=lambda x: x['final_score'], reverse=True)
        self.show_report_window(sorted_results, validated_subs)

    def show_report_window(self, sorted_results, validated_subs):
        report_win = tk.Toplevel(self.root)
        report_win.title(f"{self.settings['competition_name']} - Final Report")
        report_win.geometry("800x500")
        self.active_windows.append(report_win)
        self.update_theme()

        text_widget = tk.Text(report_win, wrap=tk.WORD, font=("Courier New", 10), padx=10, pady=10, relief=tk.FLAT)
        text_widget.pack(fill=tk.BOTH, expand=True)

        is_curved = self.settings.get('enable_curve', False)
        report_str = f"--- {self.settings['competition_name'].upper()} FINAL LEADERBOARD ---\n\n"
        
        if is_curved:
            cat_headers = [f"{name[:10]}(R/C)" for name in self.settings['category_names']]
        else:
            cat_headers = [f"{name[:15]}" for name in self.settings['category_names']]
            
        header = f"{'Rank':<5} {'Submission Name':<25} {'Final Score':<15} {' '.join(f'{h:<12}' for h in cat_headers)}\n"
        report_str += header
        report_str += "-" * len(header) + "\n"
        
        for i, res in enumerate(sorted_results, 1):
            scores_str_parts = []
            for cat_idx in range(len(res['raw_scores'])):
                raw_score = res['raw_scores'][cat_idx]
                final_cat_score = res['final_category_scores'][cat_idx]
                if is_curved:
                    scores_str_parts.append(f"{raw_score:.1f}/{final_cat_score:.1f}")
                else:
                    scores_str_parts.append(f"{final_cat_score:.1f}")

            scores_str = ' '.join(f'{s:<12}' for s in scores_str_parts)
            line = f"{i:<5} {res['name']:<25} {res['final_score']:<15.2f} {scores_str}\n"
            report_str += line
        
        report_str += "\n" + "-" * len(header) + "\n\n--- Category Winners (Highest Raw Score) ---\n"
        raw_max_scores = [max((sub['raw_scores'][i] for sub in validated_subs), default=0) for i in range(self.settings['num_categories'])]
        winners = [next((s['name'] for s in validated_subs if s['raw_scores'][i] == max_s), "N/A") for i, max_s in enumerate(raw_max_scores)]
        for i, winner in enumerate(winners):
            report_str += f"🏆 {self.settings['category_names'][i]}: {winner} ({raw_max_scores[i]:.1f} pts)\n"

        text_widget.insert(tk.END, report_str)
        text_widget.config(state=tk.DISABLED)
        
        export_button = ttk.Button(report_win, text="📄 Export to Word", command=lambda: self.export_report_to_word(sorted_results, raw_max_scores, winners))
        export_button.pack(fill=tk.X, padx=10, pady=10)

        if not DOCX_AVAILABLE or not MATPLOTLIB_AVAILABLE:
            export_button.config(state=tk.DISABLED)
            missing = []
            if not DOCX_AVAILABLE: missing.append("'python-docx'")
            if not MATPLOTLIB_AVAILABLE: missing.append("'matplotlib' & 'numpy'")
            tooltip_text = f"Install {', '.join(missing)} to enable Word export.\nRun: pip install python-docx matplotlib numpy"
            ToolTip(export_button, tooltip_text)


    def export_report_to_word(self, sorted_results, max_scores, winners):
        if not DOCX_AVAILABLE or not MATPLOTLIB_AVAILABLE:
            return
            
        is_curved = self.settings.get('enable_curve', False)
        chart_type = self.settings.get('chart_type', 'Bar Graph')
        category_colors = self.settings.get('category_colors', self.default_settings['category_colors'])

        try:
            doc = Document()
            # --- Title Page ---
            title = doc.add_heading(self.settings['competition_name'], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('Final Results').font.size = Pt(20)
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Report Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
            doc.add_page_break()

            # --- Methodology Section ---
            doc.add_heading('Scoring Methodology', level=1)
            p = doc.add_paragraph("The final scores were calculated based on the following settings:", style='List Bullet')
            if is_curved:
                doc.add_paragraph('Score Curving: Enabled', style='List Bullet 2')
                p = doc.add_paragraph(style='List Bullet 3')
                p.add_run("This feature scales scores within each category so the top score becomes the maximum possible score, proportionally increasing all other scores in that category.").italic = True
                p = doc.add_paragraph(style='List Bullet 3')
                p.add_run("Example: In a 1-100 scale, if the highest raw score was 80, it becomes 100. A raw score of 60 would then be scaled up to 75.").italic = True
            doc.add_paragraph(f"Custom Weights: {'Enabled' if self.settings['enable_weights'] else 'Disabled'}", style='List Bullet 2')
            
            # --- Chart Section (if Bar Graph) ---
            if chart_type == 'Bar Graph':
                doc.add_heading('Score Visualization', level=1)
                chart_path = 'scores_chart.png'
                fig, ax = plt.subplots(figsize=(10, 6), layout='constrained')
                
                score_label = "Curved Score" if is_curved else "Score"
                sub_names = [res['name'] for res in sorted_results]
                cat_names = self.settings['category_names']
                scores_by_cat = {name: [res['final_category_scores'][i] for res in sorted_results] for i, name in enumerate(cat_names)}
                x = np.arange(len(sub_names))
                width = 0.8 / len(cat_names)
                multiplier = 0
                
                for i, (attribute, measurement) in enumerate(scores_by_cat.items()):
                    offset = width * multiplier
                    rects = ax.bar(x + offset - (width * (len(cat_names)-1) / 2), measurement, width, label=attribute, color=category_colors[i])
                    ax.bar_label(rects, padding=3, fmt='%.1f', fontsize=8)
                    multiplier += 1
                
                ax.set_ylabel(score_label, fontweight='bold')
                ax.set_title('Submission Scores by Category', fontsize=14, fontweight='bold')
                ax.set_xticks(x, sub_names, rotation=45, ha="right")
                ax.legend(loc='upper right', ncols=1)
                ax.grid(axis='y', linestyle='--', alpha=0.7)
                ax.set_ylim(0, float(self.settings['score_max']) * 1.1)
                
                plt.savefig(chart_path, dpi=300)
                plt.close(fig)
                
                doc.add_paragraph("The following chart compares the final scores for each submission across all categories.")
                doc.add_picture(chart_path, width=Inches(6.5))
                os.remove(chart_path)
                doc.add_page_break()

            # --- Final Rankings Table ---
            doc.add_heading('Final Rankings', level=1)
            num_cat = self.settings['num_categories']
            num_cols = 3 + num_cat + (1 if chart_type == 'Pie Chart' else 0)
            table = doc.add_table(rows=1, cols=num_cols)
            table.style = 'Table Grid'
            table.autofit = False
            
            hdr_cells = table.rows[0].cells
            if is_curved:
                cat_headers = [f"{name} (Raw/Curved)" for name in self.settings['category_names']]
            else:
                cat_headers = self.settings['category_names']
            
            headers = ['Rank', 'Submission Name', 'Final Score'] 
            if chart_type == 'Pie Chart':
                headers.append("Score Visual")
            headers += cat_headers

            for i, header_text in enumerate(headers):
                p = hdr_cells[i].paragraphs[0]
                p.add_run(header_text).bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Loop through results to populate table
            max_possible_score = float(self.settings.get('score_max', 100))
            for i, res in enumerate(sorted_results, 1):
                row_cells = table.add_row().cells
                if i % 2 == 0:
                    for cell in row_cells:
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), 'E7E6E6')
                        cell._tc.get_or_add_tcPr().append(shading_elm)

                # Standard columns
                col_idx = 0
                row_cells[col_idx].text, row_cells[col_idx].paragraphs[0].alignment = str(i), WD_ALIGN_PARAGRAPH.CENTER
                col_idx += 1
                row_cells[col_idx].text = res['name']
                col_idx += 1
                row_cells[col_idx].text, row_cells[col_idx].paragraphs[0].alignment = f"{res['final_score']:.2f}", WD_ALIGN_PARAGRAPH.CENTER
                col_idx += 1

                # Pie Chart column
                if chart_type == 'Pie Chart':
                    chart_path = f"temp_pie_{i}.png"
                    fig, ax = plt.subplots(figsize=(1, 1))
                    
                    score = res['final_score']
                    # Ensure score is not negative for visualization
                    achieved_score = max(0, score)
                    missing_score = max(0, max_possible_score - achieved_score)
                    
                    pie_values = [achieved_score, missing_score]
                    pie_colors = ['#2ca02c', '#000000']
                    pie_labels = [f"{score:.1f}", None]
                    
                    ax.pie(pie_values, labels=pie_labels, colors=pie_colors, startangle=90, counterclock=False, 
                           wedgeprops={'edgecolor': 'white'}, textprops={'fontsize': 14, 'color': 'white', 'fontweight': 'bold'})
                    ax.axis('equal')
                    
                    plt.savefig(chart_path, dpi=150, transparent=True)
                    plt.close(fig)
                    
                    # Add picture to cell
                    cell = row_cells[col_idx]
                    cell.text = '' # Clear cell text
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(chart_path, width=Inches(0.65))
                    os.remove(chart_path)
                    col_idx += 1

                # Category score columns
                for j in range(num_cat):
                    if is_curved:
                        cell_text = f"{res['raw_scores'][j]:.1f} / {res['final_category_scores'][j]:.1f}"
                    else:
                        cell_text = f"{res['final_category_scores'][j]:.1f}"
                    row_cells[col_idx+j].text = cell_text
                    row_cells[col_idx+j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # --- Category Winners Section ---
            doc.add_page_break()
            doc.add_heading('Category Winners (Highest Raw Score)', level=1)
            for i, winner in enumerate(winners):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{self.settings['category_names'][i]}: ").bold = True
                p.add_run(f"{winner} ({max_scores[i]:.1f} pts)")
            
            # --- Footer ---
            section = doc.sections[0]
            footer = section.footer
            p = footer.paragraphs[0]
            p.text = f"{self.settings['competition_name']} | Page "
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            safe_filename = "".join(c for c in self.settings['competition_name'] if c.isalnum() or c in (' ', '_')).rstrip().replace(' ', '_')
            filename = f"{safe_filename}_Results_{datetime.datetime.now().strftime('%Y-%m-%d')}.docx"
            doc.save(filename)
            messagebox.showinfo("Export Successful", f"Report successfully saved as:\n{os.path.abspath(filename)}", parent=self.root)
        except Exception as e:
            messagebox.showerror("Export Error", f"An error occurred while saving the file:\n{e}", parent=self.root)

if __name__ == "__main__":
    root = tk.Tk()
    app = ScoreCalculatorApp(root)
    root.mainloop()